
from __future__ import annotations
from typing import Dict
import pandas as pd
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

def export_docx(path:str, metadata:Dict, tabla:pd.DataFrame, figs:Dict[str,str]):
    doc = Document()
    doc.add_heading(metadata.get("titulo","Informe AIP"), 0)
    doc.add_paragraph(metadata.get("resumen",""))
    doc.add_heading("Resultados", level=1)
    doc.add_paragraph(f"AIP acumulado: {metadata.get('AIP_total','')}")
    doc.add_paragraph(f"SPF final: {metadata.get('SPF_final','')}")
    doc.add_heading("Tabla principal", level=2)
    t = doc.add_table(rows=1, cols=len(tabla.columns))
    hdr = t.rows[0].cells
    for i,c in enumerate(tabla.columns): hdr[i].text = str(c)
    for _,row in tabla.iterrows():
        cells = t.add_row().cells
        for i,c in enumerate(tabla.columns):
            cells[i].text = str(row[c])
    for k,fp in figs.items():
        doc.add_heading(k, level=2)
        doc.add_picture(fp, width=Inches(6))
    doc.save(path)

def export_pdf(path:str, metadata:Dict):
    c = canvas.Canvas(path, pagesize=A4)
    w,h = A4; y = h-72
    c.setFont("Helvetica-Bold", 14)
    c.drawString(72,y, metadata.get("titulo","Informe AIP")); y -= 24
    c.setFont("Helvetica", 11)
    for line in metadata.get("resumen","").split("\n"):
        c.drawString(72,y, line[:100]); y -= 14
    y -= 8
    c.drawString(72,y, f"AIP acumulado: {metadata.get('AIP_total','')}"); y -= 14
    c.drawString(72,y, f"SPF final: {metadata.get('SPF_final','')}")
    c.showPage(); c.save()
